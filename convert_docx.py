import os
from docx import Document

def convert_docx_to_md(docx_path):
    try:
        doc = Document(docx_path)
        md_content = ""
        for para in doc.paragraphs:
            md_content += para.text + "\n\n"
        
        # Simple table extraction
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                md_content += "| " + " | ".join(row_text) + " |\n"
            md_content += "\n"

        md_path = docx_path.replace(".docx", ".md")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(md_content)
        print(f"Converted: {docx_path} -> {md_path}")
    except Exception as e:
        print(f"Error converting {docx_path}: {e}")

def main():
    base_dir = os.getcwd()
    for root, dirs, files in os.walk(base_dir):
        # Skip the 'fase de preparação' folder to avoid converting legacy stuff if not needed, 
        # but the user might want everything. Let's stick to the root and sugestoes_banca for now 
        # as that's where the relevant files are.
        # Actually, let's just convert everything that is a docx.
        for file in files:
            if file.endswith(".docx") and not file.startswith("~$"):
                full_path = os.path.join(root, file)
                convert_docx_to_md(full_path)

if __name__ == "__main__":
    main()
